import io
import os
from typing import BinaryIO

from docx import Document
from langchain_core.documents import Document as LangChainDocument
from pypdf import PdfReader

from config import validate_extension


class DocumentLoadError(Exception):
    """Raised when a document cannot be loaded or yields no text."""


def _extract_pdf_text(file_obj: BinaryIO) -> str:
    reader = PdfReader(file_obj)
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text.strip())
    return "\n\n".join(text for text in pages if text)


def _extract_txt_text(file_obj: BinaryIO) -> str:
    raw = file_obj.read()
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError("Could not decode text file.")


def _extract_docx_text(file_obj: BinaryIO) -> str:
    document = Document(file_obj)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    return "\n\n".join(text for text in paragraphs if text)


def _extract_text(filename: str, file_obj: BinaryIO) -> str:
    _, ext = os.path.splitext(filename.lower())

    if ext == ".pdf":
        return _extract_pdf_text(file_obj)
    if ext == ".txt":
        return _extract_txt_text(file_obj)
    if ext == ".docx":
        return _extract_docx_text(file_obj)

    raise DocumentLoadError(f"Unsupported file type: {ext}")


def load_uploaded_files(uploaded_files) -> list[LangChainDocument]:
    """
    Extract text from uploaded Streamlit file objects and return LangChain documents.

    Each file becomes one or more documents with source metadata.
    """
    if not uploaded_files:
        raise DocumentLoadError("No files were uploaded.")

    documents: list[LangChainDocument] = []

    for uploaded_file in uploaded_files:
        filename = uploaded_file.name

        if not validate_extension(filename):
            raise DocumentLoadError(f"Unsupported file type: {filename}")

        file_bytes = uploaded_file.getvalue()
        if not file_bytes:
            raise DocumentLoadError(f"File is empty: {filename}")

        try:
            text = _extract_text(filename, io.BytesIO(file_bytes))
        except DocumentLoadError:
            raise
        except Exception as exc:
            raise DocumentLoadError(f"Failed to read {filename}: {exc}") from exc

        if not text.strip():
            raise DocumentLoadError(f"No text could be extracted from {filename}.")

        documents.append(
            LangChainDocument(
                page_content=text,
                metadata={"source": filename},
            )
        )

    return documents
